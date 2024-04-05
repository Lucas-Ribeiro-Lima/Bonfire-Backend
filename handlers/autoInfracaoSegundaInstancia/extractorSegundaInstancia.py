import re
from docx import Document
from Classes.AutoSegundaInstancia import SegundaInstancia
from Classes import Conversores
from Exceptions.CustomExceptions import ErrDataPubli

def parseDocx(docx):

    doc = Document(docx)
    autoSegundaInstanciaList = []
    count = 0

    #Extraindo data da publicação
    data_publicacao_extracted = ""
    for paragraph in doc.paragraphs:
        data_publicacao_extracted += paragraph.text + " "

    padrao_data = r'PUBLICADO NO DIÁRIO OFICIAL DO MUNICIPIO DE BELO HORIZONTE EM (\d{2}/\d{2}/\d{4})'
    match_data_publicacao = re.search(padrao_data, data_publicacao_extracted)

    DAT_PUBL = match_data_publicacao.group(1) if match_data_publicacao else None

    if DAT_PUBL != None:
        DAT_PUBL = Conversores.Conversores.converte_data(DAT_PUBL)
    else:
        raise ErrDataPubli("Data de publicação não encontrada no documento")

    #Extraindo Numero da ata
    num_ata_extracted = ""
    num_atas = []
    padrao_num_ata = r'ATA DA (\d+)ª'
    for paragraph in doc.paragraphs:
        match_num_ata = re.search(padrao_num_ata, paragraph.text)
        if match_num_ata:
            num_atas.append(match_num_ata.group(1))

    for index, table in enumerate(doc.tables):
        for row in table.rows:
            rowData = [cell.text.strip() for cell in row.cells]

            NUM_RECURSO = rowData[0]
            NUM_AI = rowData[1]
            NOM_CONC = rowData[2]
            valida_resultado = str(rowData[3]).upper()
            if valida_resultado == 'IMPROCEDENTE':
                RESULTADO = False
            else:
                RESULTADO = True 
            NUM_ATA = num_atas[index]

            if NUM_RECURSO =='RECURSO':
                continue

            autoSegundaInstancia = SegundaInstancia(NUM_ATA, NUM_RECURSO, NUM_AI, NOM_CONC, RESULTADO, DAT_PUBL)  
            autoSegundaInstanciaList.append(autoSegundaInstancia.toDict())  
    
    return autoSegundaInstanciaList