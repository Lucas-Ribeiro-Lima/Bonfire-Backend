import re
from docx import Document
from sqlalchemy import text
from classes import *
from classes.Recurso import Recurso
from repositories.database import MySQL
from handlers import log
from exceptions.CustomExceptions import ErrDataPubli, ErrGetData, ErrInsertData, ErrNullInsert, ErrQuantityOfAtas

engine = MySQL().get_connection()

def getPrimeiraInstancia(date, ata):
    """Retorna os recursos de primeira instância"""
    try:
        with engine.connect() as conn:
            query = '''
                SELECT 
                    re.NUM_AI, 
                    re.NUM_ATA, 
                    re.DAT_PUBL,
                    ai.COD_LINH,
                    ai.NUM_VEIC,
                    ai.IDN_PLAC_VEIC
                FROM recurso_primeira_instancia re
                JOIN auto_infracao ai
                    ON re.NUM_AI = ai.NUM_AI
                WHERE 1= 1
            '''
            params = {}
            if ata is not None:
                query += " AND re.NUM_ATA = :ata"
                params["ata"] = ata
            if date is not None:
                query += " AND re.DAT_PUBL = :date"
                params["date"] = date

            query += " LIMIT 300"

            result = conn.execute(text(query), params)

            rows = result.mappings().all()
            json_data = [dict(row) for row in rows]

        engine.dispose()
        return json_data

    except Exception as e:
        log.HandleErrorLog(e)
        raise ErrGetData("Erro ao recuperar os autos de infração de segunda instancia", 500)
    
def parseDocx(docx):
    """Realiza o parse do documento DOCX de resultado de recurso e extrai os KVP e retorna uma lista de recursos"""
    doc = Document(docx)
    recurso_primeira_instancia_list = []
    
    #Extraindo data da publicação
    data_publicacao_extracted = ""
    for paragraph in doc.paragraphs:
        data_publicacao_extracted += paragraph.text + " "

    padrao_data = r'PUBLICADO NO DIÁRIO OFICIAL DO MUNICIPIO DE BELO HORIZONTE EM (\d{2}/\d{2}/\d{4})'
    match_data_publicacao = re.search(padrao_data, data_publicacao_extracted)

    DAT_PUBL = match_data_publicacao.group(1) if match_data_publicacao else None

    if DAT_PUBL != None:
        DAT_PUBL = Conversores.Conversores.converte_data(DAT_PUBL)
    else:
        raise ErrDataPubli("Data de publicação não encontrada no documento", 400)

    #Extraindo Numero da ata
    num_atas = []
    padrao_num_ata = r'ATA\s+DA\s+(\d+)ª'
    for paragraph in doc.paragraphs:
        match_num_ata = re.search(padrao_num_ata, paragraph.text)
        if match_num_ata:
            num_atas.append(match_num_ata.group(1))

    qtd_atas = len(num_atas)
    qtd_tables = len(doc.tables)
    if qtd_atas != qtd_tables:
        raise ErrQuantityOfAtas("Quantidade de atas encontradas difere da quantidade de tabelas", qtd_atas, qtd_tables, 400)
    

    for index, table in enumerate(doc.tables):
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]

            NUM_RECURSO = row_data[0]
            NUM_AI = row_data[1]
            NOM_CONC = row_data[2]
            valida_resultado = str(row_data[3]).upper()
            if valida_resultado == 'IMPROCEDENTE':
                RESULTADO = False
            else:
                RESULTADO = True 
            NUM_ATA = num_atas[index]

            if NUM_RECURSO =='RECURSO':
                continue

            recurso_primeira_instancia = Recurso(NUM_ATA, NUM_RECURSO, NUM_AI, NOM_CONC, RESULTADO, DAT_PUBL)
            recurso_primeira_instancia_list.append(recurso_primeira_instancia.toDict())
    return recurso_primeira_instancia_list


def insertPrimeiraInstancia(recursos_primeira_instancia):
    """Insere no banco de dados uma lista de recursos de primeira instância"""
    counter = 0
    query = ("INSERT IGNORE INTO recurso_primeira_instancia"
             "(NUM_AI, NUM_ATA, NUM_RECURSO, NOM_CONC, RESULTADO, DAT_PUBL)"
             " VALUES (:NUM_AI, :NUM_ATA, :NUM_RECURSO, :NOM_CONC, :RESULTADO, :DAT_PUBL)")
    
    if recursos_primeira_instancia is None:
        raise ErrNullInsert('Lista de auto vazio, nenhum registro inserido', 400)
    
    try:
        with engine.connect() as conn:
            for recurso in recursos_primeira_instancia:
                result = conn.execute(text(query), recurso)
                if result.rowcount > 0:
                    counter = counter +1                   
            conn.commit()
        engine.dispose()
        return counter
    except Exception as e:
        log.HandleErrorLog(e)

        raise ErrInsertData("Erro ao inserir os autos de segunda instância do banco", 500)
        