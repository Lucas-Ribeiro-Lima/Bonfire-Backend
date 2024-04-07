import re
import pandas as pd
from docx import Document
from sqlalchemy import text
from classes import *
from classes.AutoSegundaInstancia import SegundaInstancia
from repositories import database
from handlers import log
from exceptions.CustomExceptions import ErrDataPubli, ErrGetData, ErrInsertData, ErrNullInsert

def getSegundaInstancia():
    """Retorna os autos de infração que estão em segunda instância"""
    engine = database.mySQL().createDatabaseStringConnection()
    return_fields = "ai.NUM_AI, ai.NUM_VEIC, ai.IDN_PLAC_VEIC, ai.COD_LINH"
    query = f'''SELECT {return_fields} FROM auto_infracao ai WHERE ai.NUM_AI IN (SELECT NUM_AI FROM segundaInstancia);
    '''        
    try:
        with engine.connect():
            dataFrame = pd.read_sql(query, engine)
            jsonData = dataFrame.to_json(orient='records')           
        engine.dispose()
        return jsonData    
    except Exception as e:
        log.HandleErrorLog(e)
        raise ErrGetData("Erro ao recuperar os autos de infração de segunda instancia", 500)
    
def parseDocx(docx):
    """Realiza o parse do documento DOCX e extrai os KVP e retona uma lista com as infrações de segunda instancia"""
    doc = Document(docx)
    autoSegundaInstanciaList = []
    
    #Extraindo data da publicação
    data_publicacao_extracted = ""
    for paragraph in doc.paragraphs:
        data_publicacao_extracted += paragraph.text + " "

    padrao_data = r'PUBLICADO NO DIÁRIO OFICIAL DO MUNICIPIO DE BELO HORIZONTE EM (\d{2}/\d{2}/\d{4})'
    match_data_publicacao = re.search(padrao_data, data_publicacao_extracted)

    DAT_PUBL = match_data_publicacao.group(1) if match_data_publicacao else None

    if DAT_PUBL != None:
        DAT_PUBL = Conversores.Conversores.converte_data(DAT_PUBL)
    else:
        raise ErrDataPubli("Data de publicação não encontrada no documento", 400)

    #Extraindo Numero da ata
    num_ata_extracted = ""
    num_atas = []
    padrao_num_ata = r'ATA DA (\d+)ª'
    for paragraph in doc.paragraphs:
        match_num_ata = re.search(padrao_num_ata, paragraph.text)
        if match_num_ata:
            num_atas.append(match_num_ata.group(1))

    for index, table in enumerate(doc.tables):
        for row in table.rows:
            rowData = [cell.text.strip() for cell in row.cells]

            NUM_RECURSO = rowData[0]
            NUM_AI = rowData[1]
            NOM_CONC = rowData[2]
            valida_resultado = str(rowData[3]).upper()
            if valida_resultado == 'IMPROCEDENTE':
                RESULTADO = False
            else:
                RESULTADO = True 
            NUM_ATA = num_atas[index]

            if NUM_RECURSO =='RECURSO':
                continue

            autoSegundaInstancia = SegundaInstancia(NUM_ATA, NUM_RECURSO, NUM_AI, NOM_CONC, RESULTADO, DAT_PUBL)  
            autoSegundaInstanciaList.append(autoSegundaInstancia.toDict())  
    return autoSegundaInstanciaList


def insertSegundaInstancia(autoSegundaInstanciaList) -> int:
    """Insere no banco de dados uma lista de autos de infração de segunda instância"""
    engine = database.mySQL().createDatabaseStringConnection()
    counter = 0
    query = "INSERT IGNORE INTO segundaInstancia (NUM_AI, NUM_ATA, NUM_RECURSO, NOM_CONC, RESULTADO, DAT_PUBL) VALUES (:NUM_AI, :NUM_ATA, :NUM_RECURSO, :NOM_CONC, :RESULTADO, :DAT_PUBL)"
    
    if autoSegundaInstanciaList == None:
        raise ErrNullInsert('Lista de auto vazio, nenhum registro inserido', 400)
    
    try:
        with engine.connect() as conn:
            for autoSegundaInstancia in autoSegundaInstanciaList:
                result = conn.execute(text(query), autoSegundaInstancia)
                if result.rowcount > 0:
                    counter = counter +1                   
            conn.commit()
        engine.dispose()
        return counter
    except Exception as e:
        log.HandleErrorLog(e)

        raise ErrInsertData("Erro ao inserir os autos de segunda instância do banco", 500)
        